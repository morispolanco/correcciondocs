import streamlit as st
import docx2txt
from googletrans import Translator
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import tempfile
import os

# Título de la aplicación
st.title("Traductor de Documentos")

# Cargar archivo DOCX
uploaded_file = st.file_uploader("Cargar archivo DOCX", type=["docx"])

# Traducción al inglés con Google Translate
def translate_to_english(text):
    translator = Translator()
    translation = translator.translate(text, dest='en')
    return translation.text

# Traducción al español con Google Translate
def translate_to_spanish(text):
    translator = Translator()
    translation = translator.translate(text, dest='es')
    return translation.text

# Comparar documentos y resaltar cambios
def compare_documents(original_doc, translated_doc):
    original_text = docx2txt.process(original_doc)
    translated_text = docx2txt.process(translated_doc)

    original_lines = original_text.splitlines()
    translated_lines = translated_text.splitlines()

    compared_lines = []
    for i in range(len(original_lines)):
        original_line = original_lines[i]
        translated_line = translated_lines[i]

        if original_line != translated_line:
            compared_lines.append((original_line, translated_line))
        else:
            compared_lines.append((original_line, original_line))

    return compared_lines

# Mostrar cambios en el documento
def show_changes(compared_lines):
    doc = Document()
    for original_line, translated_line in compared_lines:
        paragraph = doc.add_paragraph()
        if original_line == translated_line:
            paragraph.add_run(original_line)
        else:
            run = paragraph.add_run(translated_line)
            run.font.highlight_color = 6

    return doc

# Botón para realizar las traducciones y comparar documentos
if st.button("Traducir y Comparar"):
    if uploaded_file is not None:
        # Leer el contenido del archivo DOCX
        original_doc = Document(uploaded_file)
        original_text = docx2txt.process(uploaded_file)

        # Traducción al inglés
        translated_text_english = translate_to_english(original_text)
        translated_doc_english = Document()
        translated_doc_english.add_paragraph(translated_text_english)

        # Traducción al español
        translated_text_spanish = translate_to_spanish(original_text)
        translated_doc_spanish = Document()
        translated_doc_spanish.add_paragraph(translated_text_spanish)

        # Guardar archivo DOCX traducido en disco
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            translated_doc_spanish.save(temp_file.name)
            translated_doc_path = temp_file.name

        # Comparar documentos
        compared_lines = compare_documents(uploaded_file, translated_doc_path)

        # Mostrar cambios
        changed_doc = show_changes(compared_lines)

        # Guardar documento con cambios en disco
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            changed_doc.save(temp_file.name)
            changed_doc_path = temp_file.name

        # Descargar documento con cambios
        st.download_button("Descargar Documento con Cambios", data=changed_doc_path, file_name="documento_con_cambios.docx")

        # Eliminar archivos temporales
        os.remove(translated_doc_path)
        os.remove(changed_doc_path)

    else:
        st.error("Por favor, cargue un archivo DOCX.")
